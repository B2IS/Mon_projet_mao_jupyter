"""
agent-tools/main.py — Serveur FastAPI d'outils multi-formats pour SIGEPP-DPE

Endpoints :
  POST /parse     — Analyse un fichier uploadé (multi-format)
  GET  /health    — Santé du service

Formats supportés :
  PDF, Excel/CSV, Word, ZIP/RAR, DXF/DWG (CAO), KML/KMZ, SHP (SIG),
  SCD/CID/ICD (SCADA CEI 61850), XER/MPP (Primavera/MS Project)
"""

import os, io, zipfile, json, traceback
from typing import Optional
from fastapi import FastAPI, UploadFile, File, HTTPException
from fastapi.middleware.cors import CORSMiddleware

app = FastAPI(title="SIGEPP-DPE Document Parser", version="1.0")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)


# ─────────────────────────────────────────────────────────────────────────────
# Outils de parsing par format
# ─────────────────────────────────────────────────────────────────────────────

def parse_pdf(data: bytes, filename: str) -> dict:
    try:
        from pypdf import PdfReader
        reader = PdfReader(io.BytesIO(data))
        text = "\n".join(p.extract_text() or "" for p in reader.pages[:20])
        return {
            "type": "pdf",
            "pages": len(reader.pages),
            "text_extract": text[:8000],
            "metadata": dict(reader.metadata or {}),
        }
    except Exception as e:
        return {"type": "pdf", "error": str(e), "text_extract": ""}


def parse_excel(data: bytes, filename: str) -> dict:
    try:
        import pandas as pd
        xl = pd.ExcelFile(io.BytesIO(data))
        sheets = {}
        for sn in xl.sheet_names[:5]:
            df = xl.parse(sn, nrows=100)
            sheets[sn] = df.to_dict(orient="records")
        return {"type": "excel", "sheets": list(xl.sheet_names), "data": sheets}
    except Exception as e:
        return {"type": "excel", "error": str(e)}


def parse_csv(data: bytes, filename: str) -> dict:
    try:
        import pandas as pd
        df = pd.read_csv(io.BytesIO(data), nrows=200, encoding="utf-8", errors="replace")
        return {
            "type": "csv",
            "columns": list(df.columns),
            "rows": len(df),
            "sample": df.head(20).to_dict(orient="records"),
        }
    except Exception as e:
        return {"type": "csv", "error": str(e)}


def parse_word(data: bytes, filename: str) -> dict:
    try:
        from docx import Document
        doc = Document(io.BytesIO(data))
        text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        tables = []
        for t in doc.tables[:5]:
            rows = [[c.text for c in r.cells] for r in t.rows[:10]]
            tables.append(rows)
        return {"type": "word", "text_extract": text[:8000], "tables": tables}
    except Exception as e:
        return {"type": "word", "error": str(e)}


def parse_dxf(data: bytes, filename: str) -> dict:
    """CAO/Plans électriques — extrait couches, textes, blocs (sous-stations, pylônes)."""
    try:
        import ezdxf
        doc = ezdxf.read(io.BytesIO(data))
        msp = doc.modelspace()

        couches = [{"nom": l.dxf.name, "couleur": l.dxf.color} for l in doc.layers]
        textes = []
        blocs = []
        for e in msp:
            if e.dxftype() == "TEXT":
                t = e.dxf.text.strip()
                if t: textes.append(t)
            elif e.dxftype() == "MTEXT":
                t = e.dxf.text.strip()
                if t: textes.append(t)
            elif e.dxftype() == "INSERT":
                blocs.append(e.dxf.name)

        # Détection du domaine électrique dans les couches
        domaine_hints = []
        couches_noms = [c["nom"].upper() for c in couches]
        if any("HTA" in n or "HTB" in n for n in couches_noms): domaine_hints.append("Réseau HTA/HTB")
        if any("BT" in n or "BTA" in n for n in couches_noms):  domaine_hints.append("Réseau BT")
        if any("POSTE" in n or "POST" in n for n in couches_noms): domaine_hints.append("Poste de transformation")
        if any("CABLE" in n or "CÂBLE" in n for n in couches_noms): domaine_hints.append("Plan câbles")
        if any("TOPO" in n or "TERRAIN" in n for n in couches_noms): domaine_hints.append("Topographie/emprise")

        return {
            "type": "dxf",
            "domaines_detectes": domaine_hints,
            "couches": couches[:30],
            "textes_extraits": textes[:50],
            "blocs_inseres": list(set(blocs))[:30],
            "nb_entites": len(list(msp)),
        }
    except Exception as e:
        return {"type": "dxf", "error": str(e)}


def parse_kml(data: bytes, filename: str) -> dict:
    """Tracés géospatiaux KML/KMZ — lignes HTA/HTB, position pylônes, emprises postes."""
    try:
        # KMZ est un ZIP contenant un KML
        if filename.lower().endswith(".kmz"):
            with zipfile.ZipFile(io.BytesIO(data)) as z:
                kml_name = next((n for n in z.namelist() if n.endswith(".kml")), None)
                if kml_name:
                    data = z.read(kml_name)

        import xml.etree.ElementTree as ET
        root = ET.fromstring(data)
        ns = {"kml": "http://www.opengis.net/kml/2.2"}

        placemarks = root.findall(".//kml:Placemark", ns) or root.findall(".//Placemark")
        features = []
        for pm in placemarks[:50]:
            name_el = pm.find("kml:name", ns) or pm.find("name")
            desc_el  = pm.find("kml:description", ns) or pm.find("description")
            coord_el = pm.find(".//kml:coordinates", ns) or pm.find(".//coordinates")
            features.append({
                "nom": name_el.text if name_el is not None else "?",
                "description": (desc_el.text or "")[:200] if desc_el is not None else "",
                "nb_coords": len((coord_el.text or "").split()) if coord_el is not None else 0,
            })

        return {
            "type": "kml",
            "nb_elements": len(placemarks),
            "elements": features,
            "resume": f"{len(placemarks)} éléments géographiques (pylônes/postes/tracés).",
        }
    except Exception as e:
        return {"type": "kml", "error": str(e)}


def parse_shapefile(data: bytes, filename: str) -> dict:
    """Shapefiles SIG — emprises centrales, postes HTB, parcelles foncières."""
    try:
        import shapefile
        sf = shapefile.Reader(shp=io.BytesIO(data))
        champs = [f[0] for f in sf.fields[1:]]
        formes = sf.shapes()
        enregistrements = sf.records()
        sample = [
            dict(zip(champs, rec)) for rec in enregistrements[:10]
        ]
        return {
            "type": "shapefile",
            "type_geometrie": sf.shapeTypeName,
            "nb_entites": len(formes),
            "champs": champs,
            "echantillon": sample,
        }
    except Exception as e:
        return {"type": "shapefile", "error": str(e)}


def parse_scd(data: bytes, filename: str) -> dict:
    """SCADA CEI 61850 — SCD/CID/ICD — configuration protections et IED de postes."""
    try:
        import xml.etree.ElementTree as ET
        root = ET.fromstring(data)

        # Namespace CEI 61850
        ns_scl = "{http://www.iec.ch/61850/2003/SCL}"

        ieds = []
        for ied in root.iter(f"{ns_scl}IED"):
            ieds.append({
                "nom": ied.get("name", "?"),
                "fabricant": ied.get("manufacturer", ""),
                "type": ied.get("type", ""),
            })
        if not ieds:
            # Essai sans namespace
            for ied in root.iter("IED"):
                ieds.append({"nom": ied.get("name", "?"), "fabricant": ied.get("manufacturer", ""), "type": ied.get("type", "")})

        voltage_levels = [vl.get("name", "?") for vl in root.iter(f"{ns_scl}VoltageLevel")]
        if not voltage_levels:
            voltage_levels = [vl.get("name", "?") for vl in root.iter("VoltageLevel")]

        return {
            "type": "scd_iec61850",
            "nb_ied": len(ieds),
            "ieds": ieds[:20],
            "niveaux_tension": voltage_levels[:10],
            "resume": f"Configuration poste CEI 61850 : {len(ieds)} IED(s) de protection.",
        }
    except Exception as e:
        return {"type": "scd_iec61850", "error": str(e)}


def parse_xer(data: bytes, filename: str) -> dict:
    """Oracle Primavera XER — extraire tâches, jalons, ressources."""
    try:
        text = data.decode("utf-8", errors="replace")
        lines = text.splitlines()

        tables = {}
        current_table = None
        current_fields: list[str] = []
        current_rows: list[dict] = []

        for line in lines:
            if line.startswith("%T\t"):
                if current_table and current_rows:
                    tables[current_table] = current_rows[:200]
                current_table = line[3:].strip()
                current_fields = []
                current_rows = []
            elif line.startswith("%F\t"):
                current_fields = line[3:].strip().split("\t")
            elif line.startswith("%R\t") and current_fields:
                vals = line[3:].strip().split("\t")
                row = dict(zip(current_fields, vals))
                current_rows.append(row)

        if current_table and current_rows:
            tables[current_table] = current_rows[:200]

        taches = tables.get("TASK", [])
        jalons = [t for t in taches if t.get("task_type", "") == "TT_Mile"]
        ressources = tables.get("RSRC", [])

        return {
            "type": "primavera_xer",
            "tables_trouvees": list(tables.keys()),
            "nb_taches": len(taches),
            "nb_jalons": len(jalons),
            "nb_ressources": len(ressources),
            "taches_echantillon": taches[:10],
            "jalons": jalons[:20],
        }
    except Exception as e:
        return {"type": "primavera_xer", "error": str(e)}


def parse_msp_xml(data: bytes, filename: str) -> dict:
    """MS Project XML — extraire tâches, ressources, calendrier."""
    try:
        import xml.etree.ElementTree as ET
        root = ET.fromstring(data)
        ns = {"ms": "http://schemas.microsoft.com/project"}

        taches = []
        for t in root.findall(".//ms:Task", ns):
            nom_el = t.find("ms:Name", ns)
            dur_el = t.find("ms:Duration", ns)
            start_el = t.find("ms:Start", ns)
            finish_el = t.find("ms:Finish", ns)
            milestone_el = t.find("ms:Milestone", ns)
            taches.append({
                "nom": nom_el.text if nom_el is not None else "?",
                "duree": dur_el.text if dur_el is not None else "?",
                "debut": start_el.text if start_el is not None else "?",
                "fin": finish_el.text if finish_el is not None else "?",
                "jalon": milestone_el is not None and milestone_el.text == "1",
            })

        return {
            "type": "msproject_xml",
            "nb_taches": len(taches),
            "taches_echantillon": taches[:15],
            "jalons": [t for t in taches if t["jalon"]][:10],
        }
    except Exception as e:
        return {"type": "msproject_xml", "error": str(e)}


def parse_zip(data: bytes, filename: str) -> dict:
    """Archive ZIP — liste le contenu et analyse récursivement."""
    try:
        with zipfile.ZipFile(io.BytesIO(data)) as z:
            membres = z.namelist()
            analyses = []
            for m in membres[:20]:
                ext = m.rsplit(".", 1)[-1].lower() if "." in m else ""
                size = z.getinfo(m).file_size
                analyses.append({"chemin": m, "extension": ext, "taille": size})
            return {
                "type": "zip",
                "nb_fichiers": len(membres),
                "contenu": analyses,
                "resume": f"Archive avec {len(membres)} fichiers. Formats: {list(set(a['extension'] for a in analyses))}",
            }
    except Exception as e:
        return {"type": "zip", "error": str(e)}


# ─────────────────────────────────────────────────────────────────────────────
# Routeur principal
# ─────────────────────────────────────────────────────────────────────────────

PARSERS = {
    "pdf":   parse_pdf,
    "xlsx":  parse_excel,
    "xls":   parse_excel,
    "csv":   parse_csv,
    "docx":  parse_word,
    "doc":   parse_word,
    "dxf":   parse_dxf,
    "dwg":   lambda d, f: {"type": "dwg", "info": "DWG natif — convertir en DXF avec LibreCAD/ODA File Converter. DXF recommandé."},
    "kml":   parse_kml,
    "kmz":   parse_kml,
    "shp":   parse_shapefile,
    "scd":   parse_scd,
    "cid":   parse_scd,
    "icd":   parse_scd,
    "xer":   parse_xer,
    "mpp":   lambda d, f: {"type": "mpp", "info": "MPP natif — utiliser l'export XML de MS Project. XER ou XML recommandé."},
    "xml":   parse_msp_xml,
    "zip":   parse_zip,
}


@app.post("/parse")
async def parse_document(file: UploadFile = File(...)):
    """Analyse un document et retourne ses données structurées."""
    data = await file.read()
    ext = (file.filename or "").rsplit(".", 1)[-1].lower() if "." in (file.filename or "") else ""
    parser = PARSERS.get(ext)

    if not parser:
        raise HTTPException(status_code=415, detail=f"Format .{ext} non supporté. Formats: {list(PARSERS.keys())}")

    try:
        result = parser(data, file.filename or "")
        result["filename"] = file.filename
        result["size_bytes"] = len(data)
        return result
    except Exception as e:
        return {
            "filename": file.filename,
            "ext": ext,
            "error": str(e),
            "trace": traceback.format_exc()[-500:],
        }


@app.get("/health")
def health():
    return {"status": "ok", "service": "sigepp-document-parser", "version": "1.0"}


@app.get("/formats")
def formats():
    """Liste les formats supportés."""
    return {
        "formats": list(PARSERS.keys()),
        "domaines": {
            "Documents":   ["pdf", "docx", "doc"],
            "Tableurs":    ["xlsx", "xls", "csv"],
            "CAO/Plans":   ["dxf", "dwg"],
            "SIG/Géo":     ["kml", "kmz", "shp"],
            "SCADA":       ["scd", "cid", "icd"],
            "Planification": ["xer", "mpp", "xml"],
            "Archives":    ["zip"],
        }
    }
